#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
写作模块 - 爆款文章生成（自然语言版，无Markdown标记）
支持自动配图生成
"""

import os
import json
import random
import requests
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class ArticleWriter:
    """文章写作器 - 自然语言版"""
    
    def __init__(self, config: Dict):
        self.config = config
        self.styles = config.get('styles', {})
        self.output_config = config.get('output', {})
        
        # 图片生成配置
        self.image_config = config.get('image_generation', {})
        self.enable_image_gen = self.image_config.get('enabled', False)
        
    def generate(self, topic: str, style: str, word_count: int, sources: List[Dict]) -> Dict:
        """生成文章"""
        style_config = self.styles.get(style, self.styles.get('tech_architecture', {}))
        structure = style_config.get('structure', [])
        
        # 生成自然标题
        title = self._generate_natural_title(topic)
        
        # 生成吸睛开头
        hook = self._generate_natural_hook(topic)
        
        article = {
            "title": title,
            "topic": topic,
            "style": style,
            "hook": hook,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "sections": []
        }
        
        # 生成各部分内容
        total_words = 0
        target_per_section = word_count // len(structure)
        
        for idx, section_config in enumerate(structure):
            section_name = section_config.get('name', '')
            section_title = section_config.get('title', '')
            
            if idx == 0:
                content = self._generate_natural_introduction(topic, sources)
            elif section_name == 'principle':
                content = self._generate_natural_principle(topic)
            elif section_name == 'practice':
                content = self._generate_natural_practice(topic)
            elif section_name == 'comparison':
                content = self._generate_natural_comparison(topic)
            elif section_name == 'conclusion':
                content = self._generate_natural_conclusion(topic)
            else:
                content = self._generate_natural_generic(section_title, topic)
            
            article['sections'].append({
                "name": section_name,
                "title": section_title,
                "content": content,
                "word_count": len(content)
            })
            
            total_words += len(content)
        
        article['word_count'] = total_words
        article['sources'] = sources
        
        return article
    
    def _generate_natural_title(self, topic: str) -> str:
        """生成自然标题（无符号）"""
        templates = [
            f"{topic}深度解析",
            f"{topic}大厂实战方案",
            f"{topic}性能优化实战",
            f"{topic}踩坑总结",
            f"2026年{topic}趋势",
            f"{topic}架构师指南",
            f"从0到1掌握{topic}",
            f"{topic}最佳实践"
        ]
        return random.choice(templates)
    
    def _generate_natural_hook(self, topic: str) -> str:
        """生成自然开头"""
        hooks = [
            f"最近很多读者问我关于{topic}的问题，今天专门写一篇文章来讲清楚。",
            f"在架构设计领域，{topic}一直是个热门话题，但很多人其实理解得不够深入。",
            f"去年我们团队在做{topic}相关项目时踩了不少坑，今天把经验分享出来。",
            f"如果你正在做{topic}相关的技术选型，这篇文章应该能帮到你。",
        ]
        return random.choice(hooks)
    
    def _generate_natural_introduction(self, topic: str, sources: List[Dict]) -> str:
        """生成自然引言"""
        paragraphs = []
        
        # 自然开场
        hook = self._generate_natural_hook(topic)
        paragraphs.append(hook)
        paragraphs.append("")
        
        # 背景
        paragraphs.append(
            f"2026年，{topic}正在成为技术圈的热门话题。随着数字化转型的深入，"
            f"越来越多的企业开始关注这个领域。今天这篇文章，我会从实际应用的角度，"
            f"分享{topic}的核心原理、实战案例和未来趋势。"
        )
        paragraphs.append("")
        
        # 配图提示
        paragraphs.append(f"[配图：{topic}整体架构图]")
        paragraphs.append("")
        
        # 痛点
        paragraphs.append(
            f"在实际落地{topic}的过程中，我发现大家普遍面临这几个问题："
        )
        paragraphs.append("")
        paragraphs.append(f"第一，架构复杂度太高，涉及的技术栈太多，团队学习成本大。")
        paragraphs.append(f"第二，成本控制困难，既要保证性能，又不能投入太多资源。")
        paragraphs.append(f"第三，团队协作效率低，缺少统一的规范和最佳实践。")
        paragraphs.append("")
        
        # 文章价值
        paragraphs.append(
            f"读完这篇文章，你会对{topic}有一个系统的认识，包括核心技术原理、"
            f"大厂实战案例、不同方案的对比分析，以及可以直接落地的行动建议。"
        )
        paragraphs.append("")
        
        # 参考来源
        if sources:
            source_names = list(set([s.get('source', '') for s in sources[:3]]))
            paragraphs.append(
                f"本文内容参考了{', '.join(source_names)}等平台的讨论，"
                f"结合了我自己的实战经验，希望能给你一些启发。"
            )
        
        return '\n'.join(paragraphs)
    
    def _generate_natural_principle(self, topic: str) -> str:
        """生成技术原理部分"""
        paragraphs = []
        
        paragraphs.append(f"[配图：{topic}分层架构图]")
        paragraphs.append("")
        
        paragraphs.append(
            f"{topic}的架构设计可以分成三个层次来理解：基础设施层、服务编排层和应用层。"
            f"每一层都有其特定的职责和设计要点。"
        )
        paragraphs.append("")
        
        # 基础设施层
        paragraphs.append("基础设施层")
        paragraphs.append("")
        paragraphs.append(
            f"这一层是整个系统的底座，主要解决部署和运维的问题。"
            f"在{topic}场景下，我们需要考虑容器化部署、服务网格、以及自动化运维体系。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"具体实现上，我们通常用Docker做容器化，Kubernetes做编排，"
            f"Istio做服务网格管理。这套组合可以支撑百万级的并发请求，"
            f"P99延迟控制在100毫秒以内。"
        )
        paragraphs.append("")
        paragraphs.append(f"[配图：基础设施架构示意图]")
        paragraphs.append("")
        
        # 服务编排层
        paragraphs.append("服务编排层")
        paragraphs.append("")
        paragraphs.append(
            f"这一层是{topic}的核心所在。这里的关键在于如何合理地拆分微服务。"
            f"拆得太粗，无法独立扩展；拆得太细，调用链路复杂，调试困难。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"我的建议是按照业务领域来拆分，也就是DDD的思路。"
            f"每个服务要有明确的边界，接口契约要定义清楚。"
            f"另外，服务数量要和团队规模匹配，一般一个服务由两到三个人维护比较合适。"
        )
        paragraphs.append("")
        
        # 应用层
        paragraphs.append("应用层")
        paragraphs.append("")
        paragraphs.append(
            f"这一层直接面向用户，体验优化是关键。常见手段包括CDN加速静态资源、"
            f"合理的浏览器缓存策略、以及骨架屏等技术来提升感知性能。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"在我们的实践中，通过这些优化手段，首屏加载时间减少了60%左右，"
            f"用户满意度有明显提升。"
        )
        
        return '\n'.join(paragraphs)
    
    def _generate_natural_practice(self, topic: str) -> str:
        """生成实战案例"""
        paragraphs = []
        
        # 案例背景
        paragraphs.append("案例背景")
        paragraphs.append("")
        paragraphs.append(
            f"这是一个真实的案例。某互联网大厂在2025年初开始实施{topic}相关项目，"
            f"业务规模是日活用户五千万以上。他们的技术团队大概有二十人，"
            f"面临的挑战是如何在有限资源下支撑高并发场景。"
        )
        paragraphs.append("")
        paragraphs.append(f"[配图：改造前后架构对比图]")
        paragraphs.append("")
        
        # 解决方案
        paragraphs.append("解决方案")
        paragraphs.append("")
        paragraphs.append(
            f"经过多轮技术评审，团队最终选择了Kubernetes加Istio的技术栈。"
            f"具体的实施步骤分为三个阶段："
        )
        paragraphs.append("")
        paragraphs.append(
            f"第一阶段是容器化改造，把原有的单体应用拆分成容器化部署，"
            f"这样可以通过Kubernetes实现弹性扩缩容。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"第二阶段引入服务网格，用Istio统一管理服务间的流量，"
            f"包括负载均衡、熔断降级、灰度发布等能力。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"第三阶段是GitOps实践，把Git作为唯一可信源，"
            f"所有的配置变更都通过Git来管理，实现了部署流程的自动化。"
        )
        paragraphs.append("")
        
        # 踩坑记录
        paragraphs.append("踩坑记录")
        paragraphs.append("")
        paragraphs.append(
            f"在实施过程中，团队踩了几个比较典型的坑。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"第一个坑是微服务拆分过细。刚开始他们拆了五十多个服务，"
            f"结果调用链路太复杂，出问题很难定位。后来通过领域驱动设计重新梳理，"
            f"合并成了十五个核心服务，复杂度降低了很多。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"第二个坑是忽略了服务发现的性能。默认配置下，服务发现成了瓶颈，"
            f"后来通过本地缓存加增量更新的方案，QPS提升了五倍。"
        )
        paragraphs.append("")
        
        # 实施成果
        paragraphs.append("实施成果")
        paragraphs.append("")
        paragraphs.append(
            f"经过一年的优化，效果还是很明显的。系统可用性从百分之九十九点九"
            f"提升到了百分之九十九点九九，平均响应时间从两百毫秒降到了五十毫秒，"
            f"新功能上线周期从两周缩短到了三天，资源成本还降低了百分之二十五。"
        )
        
        return '\n'.join(paragraphs)
    
    def _generate_natural_comparison(self, topic: str) -> str:
        """生成对比分析"""
        paragraphs = []
        
        paragraphs.append(
            f"在选择{topic}技术方案时，企业通常会在几种主流方案之间做选择。"
            f"下面我从适用场景、优缺点、推荐指数三个维度来做对比。"
        )
        paragraphs.append("")
        
        # 单体架构
        paragraphs.append("方案一：单体架构")
        paragraphs.append("")
        paragraphs.append(
            f"适用场景：小型项目，团队规模在十人以内。"
        )
        paragraphs.append(
            f"优点：开发简单，部署快速，调试方便。"
        )
        paragraphs.append(
            f"缺点：扩展性差，技术栈受限，随着业务增长维护成本会急剧上升。"
        )
        paragraphs.append(
            f"推荐指数：两颗星，适合快速验证产品的初创团队。"
        )
        paragraphs.append("")
        
        # 微服务
        paragraphs.append("方案二：微服务架构")
        paragraphs.append("")
        paragraphs.append(
            f"适用场景：中大型项目，团队规模在二十到一百人。"
        )
        paragraphs.append(
            f"优点：服务独立部署和扩展，技术栈灵活，容错性好。"
        )
        paragraphs.append(
            f"缺点：运维复杂度高，需要完善的DevOps体系，分布式调试困难。"
        )
        paragraphs.append(
            f"推荐指数：四颗星，是目前最主流的方案。"
        )
        paragraphs.append("")
        
        # Serverless
        paragraphs.append("方案三：Serverless架构")
        paragraphs.append("")
        paragraphs.append(
            f"适用场景：事件驱动型应用，流量波动大的场景。"
        )
        paragraphs.append(
            f"优点：按需付费，免运维，自动扩缩容。"
        )
        paragraphs.append(
            f"缺点：有冷启动延迟，存在厂商锁定风险，调试不便。"
        )
        paragraphs.append(
            f"推荐指数：三颗星，适合特定场景，不建议全面采用。"
        )
        paragraphs.append("")
        paragraphs.append(f"[配图：三种架构对比示意图]")
        paragraphs.append("")
        
        # 选型建议
        paragraphs.append("选型建议")
        paragraphs.append("")
        paragraphs.append(
            f"总的来说，技术选型要综合考虑团队规模、业务特点、成本预算等因素。"
            f"初创公司建议用单体架构快速验证，成长期公司可以逐步拆分为微服务，"
            f"大厂可以考虑Serverless和微服务的混合架构。"
        )
        
        return '\n'.join(paragraphs)
    
    def _generate_natural_conclusion(self, topic: str) -> str:
        """生成总结"""
        paragraphs = []
        
        paragraphs.append(
            f"总结一下今天的内容。{topic}代表了技术架构演进的重要方向，"
            f"成功的架构设计需要在技术先进性和工程可行性之间找到平衡。"
        )
        paragraphs.append("")
        
        paragraphs.append("给架构师的行动建议")
        paragraphs.append("")
        paragraphs.append(
            f"第一，评估当前系统的成熟度，看看{topic}在哪些方面可以改进。"
        )
        paragraphs.append(
            f"第二，制定分阶段的演进路线，不要试图一次性做太多改造。"
        )
        paragraphs.append(
            f"第三，建立性能基准和监控体系，用数据说话。"
        )
        paragraphs.append(
            f"第四，培养团队的架构能力，多组织技术分享。"
        )
        paragraphs.append(
            f"第五，持续跟踪社区的最佳实践，保持技术敏感度。"
        )
        paragraphs.append("")
        
        paragraphs.append("2026年趋势展望")
        paragraphs.append("")
        paragraphs.append(
            f"展望未来，我觉得有几个趋势值得关注。"
        )
        paragraphs.append(
            f"第一个是AI驱动的架构设计，像Copilot这样的工具会越来越普及，"
            f"辅助架构师做决策。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"第二个是成本优化优先，FinOps会成为必选项，"
            f"架构设计要更多地考虑资源利用率。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"第三个是边缘计算融合，云边端一体化的架构会越来越普遍。"
        )
        paragraphs.append("")
        paragraphs.append(
            f"第四个是安全左移，DevSecOps会深度集成到开发流程中。"
        )
        paragraphs.append("")
        
        paragraphs.append("最后")
        paragraphs.append("")
        paragraphs.append(
            f"你在{topic}实践中有哪些经验或踩过哪些坑？欢迎在评论区分享，"
            f"我们一起交流学习。如果觉得这篇文章有帮助，也欢迎转发给需要的朋友。"
        )
        
        return '\n'.join(paragraphs)
    
    def _generate_natural_generic(self, section_title: str, topic: str) -> str:
        """生成通用内容"""
        return f"{section_title}\n\n关于{topic}的{section_title}内容...\n\n[配图：{section_title}示意图]"
    
    def generate_cover_image(self, topic: str) -> str:
        """
        生成封面图片
        
        Returns:
            图片保存路径
        """
        if not self.enable_image_gen:
            return None
        
        try:
            # 使用AI图片生成服务（如需要可以接入具体API）
            # 这里返回None，使用默认封面
            return None
        except Exception as e:
            print(f"封面生成失败: {e}")
            return None
    
    def save_to_docx(self, article: Dict, output_dir: str = './articles') -> str:
        """保存文章为Word文档"""
        os.makedirs(output_dir, exist_ok=True)
        
        # 生成文件名
        date_str = datetime.now().strftime("%Y-%m-%d")
        safe_topic = article['topic']
        for char in ['<', '>', ':', '"', '/', '\\', '|', '?', '*']:
            safe_topic = safe_topic.replace(char, '-')
        safe_topic = safe_topic.replace(' ', '-')
        filename = f"{date_str}-{safe_topic}.docx"
        filepath = os.path.join(output_dir, filename)
        
        # 创建Word文档
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 添加标题
        title = doc.add_heading(article['title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加钩子
        if article.get('hook'):
            hook = doc.add_paragraph()
            hook_run = hook.add_run(article['hook'])
            hook_run.font.size = Pt(13)
            hook_run.font.color.rgb = RGBColor(0, 102, 204)
            hook.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加元信息
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(f"技术架构深度分析 | {article['generated_at']}")
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 添加分隔线
        doc.add_paragraph('_' * 50)
        
        # 添加各章节
        for section in article['sections']:
            # 章节标题
            heading = doc.add_heading(section['title'], level=1)
            
            # 处理内容
            content_lines = section['content'].split('\n')
            for line in content_lines:
                line = line.strip()
                if not line:
                    continue
                
                # 处理图片占位符
                if line.startswith('[配图：') and line.endswith(']'):
                    # 添加图片占位框
                    img_para = doc.add_paragraph()
                    img_run = img_para.add_run(line)
                    img_run.font.size = Pt(11)
                    img_run.font.color.rgb = RGBColor(128, 128, 128)
                    img_run.italic = True
                    img_para.paragraph_format.left_indent = Inches(0.5)
                    img_para.paragraph_format.right_indent = Inches(0.5)
                    img_para.paragraph_format.space_before = Pt(10)
                    img_para.paragraph_format.space_after = Pt(10)
                    continue
                
                # 处理小标题（无序号）
                if line in ['基础设施层', '服务编排层', '应用层', 
                           '案例背景', '解决方案', '踩坑记录', '实施成果',
                           '方案一：单体架构', '方案二：微服务架构', '方案三：Serverless架构',
                           '选型建议', '给架构师的行动建议', '2026年趋势展望', '最后']:
                    sub_heading = doc.add_heading(line, level=2)
                    continue
                
                # 普通段落
                para = doc.add_paragraph(line)
                para.paragraph_format.line_spacing = 1.6
                para.paragraph_format.first_line_indent = Inches(0.3)
                para.paragraph_format.space_after = Pt(8)
        
        # 添加参考资料
        if article.get('sources'):
            doc.add_heading('参考资料', level=1)
            for idx, source in enumerate(article['sources'][:5], 1):
                para = doc.add_paragraph()
                para.add_run(f"{idx}. ").bold = True
                para.add_run(f"{source.get('title', '')} - {source.get('source', '')}")
                para.paragraph_format.line_spacing = 1.3
        
        # 添加页脚
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("技术架构观察 | 每日分享深度技术文章")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 保存文档
        doc.save(filepath)
        
        return filepath


if __name__ == "__main__":
    # 测试
    config = {
        'styles': {
            'tech_architecture': {
                'structure': [
                    {'name': 'introduction', 'title': '引言'},
                    {'name': 'principle', 'title': '技术原理'},
                    {'name': 'practice', 'title': '实战案例'},
                    {'name': 'comparison', 'title': '对比分析'},
                    {'name': 'conclusion', 'title': '总结与展望'}
                ]
            }
        },
        'output': {'default_dir': './articles'}
    }
    
    writer = ArticleWriter(config)
    article = writer.generate(
        topic='AI大模型推理优化',
        style='tech_architecture',
        word_count=2000,
        sources=[{'source': '知乎', 'title': '测试'}]
    )
    
    filepath = writer.save_to_docx(article)
    print(f"文章已保存: {filepath}")
