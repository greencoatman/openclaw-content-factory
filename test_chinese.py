import sys
sys.path.insert(0, 'scripts')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 测试生成简单的 Word 文档
doc = Document()

# 设置中文字体
doc.styles['Normal'].font.name = 'Microsoft YaHei'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 添加标题
title = doc.add_heading('测试标题：AI大模型推理优化', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加内容
para = doc.add_paragraph('这是一段测试内容。2026年，AI大模型推理优化正在成为技术圈最热门的话题之一。')
para.paragraph_format.line_spacing = 1.5

# 保存
test_path = './articles/test_chinese.docx'
doc.save(test_path)
print(f'✅ 文档已保存: {test_path}')

# 立即读取验证
doc2 = Document(test_path)
print(f'\n📝 读取验证:')
for i, para in enumerate(doc2.paragraphs[:3]):
    print(f'  段落{i}: {para.text}')
